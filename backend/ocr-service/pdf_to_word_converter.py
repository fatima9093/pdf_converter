#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF to Word Conversion Service
Converts PDFs to Word documents using pdf2docx for text-based PDFs
and OCR + pdf2docx for scanned PDFs
"""

import argparse
import json
import sys
import os
from pathlib import Path
from typing import List, Dict, Tuple, Union
import traceback
import tempfile

# Ensure UTF-8 encoding for output
if sys.stdout.encoding != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')
if sys.stderr.encoding != 'utf-8':
    sys.stderr.reconfigure(encoding='utf-8')

# PDF conversion
from pdf2docx import Converter

# PDF and image processing for OCR
from pdf2image import convert_from_path
import pytesseract
from PIL import Image

# Configure Tesseract path for Windows
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
import cv2
import numpy as np

# Document generation for OCR results
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class PDFToWordConverter:
    def __init__(self):
        """Initialize the PDF to Word converter"""
        print("INFO: PDF to Word Converter initialized", file=sys.stderr)

    def convert_text_based_pdf(self, pdf_path: str, output_path: str) -> bool:
        """
        Convert text-based PDF to Word using pdf2docx
        
        Args:
            pdf_path: Path to the input PDF file
            output_path: Path for the output Word document
            
        Returns:
            bool: True if conversion successful, False otherwise
        """
        try:
            print(f"INFO: Converting text-based PDF using pdf2docx: {pdf_path}", file=sys.stderr)
            
            # Create converter instance
            cv = Converter(pdf_path)
            
            # Convert PDF to Word document
            cv.convert(output_path, start=0, end=None)
            cv.close()
            
            print(f"SUCCESS: Text-based PDF converted to Word: {output_path}", file=sys.stderr)
            return True
            
        except Exception as e:
            print(f"ERROR: Failed to convert text-based PDF with pdf2docx: {e}", file=sys.stderr)
            traceback.print_exc()
            return False

    def convert_scanned_pdf(self, pdf_path: str, output_path: str, dpi: int = 300) -> bool:
        """
        Convert scanned PDF to Word using OCR + pdf2docx
        
        Args:
            pdf_path: Path to the input PDF file
            output_path: Path for the output Word document
            dpi: Resolution for PDF to image conversion
            
        Returns:
            bool: True if conversion successful, False otherwise
        """
        try:
            print(f"INFO: Converting scanned PDF using OCR + pdf2docx: {pdf_path}", file=sys.stderr)
            
            # Step 1: Extract text using OCR with post-processing
            pages_data = self._extract_text_with_ocr(pdf_path, dpi)
            
            if not pages_data:
                print("ERROR: No text could be extracted from the scanned PDF", file=sys.stderr)
                return False
            
            # Step 1.5: Post-process OCR results to improve accuracy
            pages_data = self._post_process_ocr_results(pages_data)
            
            # Step 2: Create a temporary PDF with the OCR text
            temp_pdf_path = self._create_text_pdf_from_ocr(pages_data, pdf_path)
            
            if not temp_pdf_path:
                print("ERROR: Failed to create temporary text PDF from OCR results", file=sys.stderr)
                return False
            
            try:
                # Step 3: Convert the temporary text PDF to Word using pdf2docx
                print(f"INFO: Converting OCR-generated PDF to Word using pdf2docx", file=sys.stderr)
                cv = Converter(temp_pdf_path)
                cv.convert(output_path, start=0, end=None)
                cv.close()
                
                print(f"SUCCESS: Scanned PDF converted to Word via OCR + pdf2docx: {output_path}", file=sys.stderr)
                return True
                
            finally:
                # Clean up temporary PDF
                if os.path.exists(temp_pdf_path):
                    os.remove(temp_pdf_path)
                    print(f"INFO: Cleaned up temporary PDF: {temp_pdf_path}", file=sys.stderr)
            
        except Exception as e:
            print(f"ERROR: Failed to convert scanned PDF: {e}", file=sys.stderr)
            traceback.print_exc()
            return False

    def _extract_text_with_ocr(self, pdf_path: str, dpi: int = 300) -> List[Dict]:
        """
        Extract text from PDF using OCR
        
        Args:
            pdf_path: Path to the PDF file
            dpi: Resolution for PDF to image conversion
            
        Returns:
            List of pages with extracted text and layout info
        """
        try:
            print(f"INFO: Converting PDF to images for OCR: {pdf_path}", file=sys.stderr)
            
            # Convert PDF pages to images
            pages = convert_from_path(pdf_path, dpi=dpi)
            print(f"INFO: Found {len(pages)} page(s)", file=sys.stderr)
            
            extracted_pages = []
            
            for page_num, page_image in enumerate(pages, 1):
                print(f"INFO: Processing page {page_num} with OCR...", file=sys.stderr)
                
                # Convert PIL image to numpy array for OpenCV
                page_array = np.array(page_image)
                
                # Preprocess image for better OCR
                processed_image = self._preprocess_image(page_array)
                
                # Extract text using Tesseract
                text_data = self._extract_with_tesseract(processed_image)
                
                page_info = {
                    'page_number': page_num,
                    'text_blocks': text_data,
                    'full_text': ' '.join([block['text'] for block in text_data if block['text'].strip()]),
                    'image_size': page_image.size
                }
                
                extracted_pages.append(page_info)
                print(f"SUCCESS: Page {page_num}: {len(text_data)} text blocks, {len(page_info['full_text'])} characters", file=sys.stderr)
            
            return extracted_pages
            
        except Exception as e:
            print(f"ERROR: Error extracting text with OCR: {e}", file=sys.stderr)
            traceback.print_exc()
            return []

    def _preprocess_image(self, image: np.ndarray) -> np.ndarray:
        """Ultra-advanced image preprocessing for maximum OCR accuracy"""
        # Convert to grayscale
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_RGB2GRAY)
        else:
            gray = image
        
        # Step 1: Noise reduction with bilateral filter (preserves edges better)
        denoised = cv2.bilateralFilter(gray, 9, 75, 75)
        
        # Step 2: Enhance contrast using multiple methods
        # CLAHE for local contrast enhancement
        clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8,8))
        enhanced = clahe.apply(denoised)
        
        # Gamma correction for better visibility
        gamma = 1.2
        lookup_table = np.array([((i / 255.0) ** (1.0 / gamma)) * 255 for i in np.arange(0, 256)]).astype("uint8")
        gamma_corrected = cv2.LUT(enhanced, lookup_table)
        
        # Step 3: Sharpening to make text edges clearer
        kernel_sharpen = np.array([[-1,-1,-1],
                                  [-1, 9,-1],
                                  [-1,-1,-1]])
        sharpened = cv2.filter2D(gamma_corrected, -1, kernel_sharpen)
        
        # Step 4: Try multiple thresholding methods and combine
        # Method 1: Adaptive threshold (Gaussian)
        thresh1 = cv2.adaptiveThreshold(sharpened, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 11, 2)
        
        # Method 2: Adaptive threshold (Mean)
        thresh2 = cv2.adaptiveThreshold(sharpened, 255, cv2.ADAPTIVE_THRESH_MEAN_C, cv2.THRESH_BINARY, 15, 3)
        
        # Method 3: Otsu's thresholding
        _, thresh3 = cv2.threshold(sharpened, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # Combine thresholding results (take intersection for cleaner text)
        combined_thresh = cv2.bitwise_and(thresh1, cv2.bitwise_and(thresh2, thresh3))
        
        # Step 5: Morphological operations to clean up and connect text
        # Remove small noise
        kernel_noise = cv2.getStructuringElement(cv2.MORPH_ELLIPSE, (2, 2))
        cleaned = cv2.morphologyEx(combined_thresh, cv2.MORPH_OPEN, kernel_noise)
        
        # Close small gaps in text
        kernel_close = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 1))
        final_image = cv2.morphologyEx(cleaned, cv2.MORPH_CLOSE, kernel_close)
        
        return final_image

    def _extract_with_tesseract(self, image: np.ndarray) -> List[Dict]:
        """Extract text using Tesseract OCR with enhanced layout and font detection"""
        # Configure Tesseract for maximum accuracy and layout preservation
        # PSM 3 = Fully automatic page segmentation, but no OSD
        # PSM 6 = Uniform block of text
        # Try multiple PSM modes for better results
        configs = [
            r'--oem 3 --psm 3 -c preserve_interword_spaces=1 -c textord_heavy_nr=1',
            r'--oem 3 --psm 6 -c preserve_interword_spaces=1 -c textord_heavy_nr=1',
            r'--oem 3 --psm 1 -c preserve_interword_spaces=1'
        ]
        
        best_data = None
        best_confidence = 0
        
        # Try different configurations and pick the best result
        for config in configs:
            try:
                data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT, config=config)
                avg_conf = np.mean([int(conf) for conf in data['conf'] if int(conf) > 0])
                if avg_conf > best_confidence:
                    best_confidence = avg_conf
                    best_data = data
            except:
                continue
        
        if best_data is None:
            # Fallback to basic configuration
            best_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
        
        text_blocks = []
        
        for i in range(len(best_data['text'])):
            text = best_data['text'][i].strip()
            conf = int(best_data['conf'][i])
            
            # Lower confidence threshold and include more text
            if text and conf > 20:  # Lower threshold for better text capture
                x, y, w, h = best_data['left'][i], best_data['top'][i], best_data['width'][i], best_data['height'][i]
                
                # More accurate font characteristics detection
                font_size = self._estimate_font_size_accurate(w, h, len(text), best_data, i)
                font_weight = self._detect_font_weight_improved(image, x, y, w, h)
                text_type = self._classify_text_type_improved(text, font_size, y, best_data, i)
                
                text_blocks.append({
                    'text': text,
                    'bbox': (x, y, x + w, y + h),
                    'confidence': conf,
                    'font_size': font_size,
                    'font_weight': font_weight,
                    'text_type': text_type,
                    'line_height': h,
                    'word_num': best_data['word_num'][i],
                    'line_num': best_data['line_num'][i],
                    'par_num': best_data['par_num'][i],
                    'block_num': best_data['block_num'][i]
                })
        
        return text_blocks

    def _create_text_pdf_from_ocr(self, pages_data: List[Dict], original_pdf_path: str) -> str:
        """
        Create a temporary text-based PDF from OCR results
        This PDF can then be converted using pdf2docx for better formatting
        """
        try:
            # Create temporary Word document first
            temp_dir = tempfile.mkdtemp()
            temp_docx_path = os.path.join(temp_dir, "ocr_temp.docx")
            
            # Create Word document from OCR data
            doc = Document()
            
            # Add title
            title = doc.add_heading('OCR Converted Document', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for page_data in pages_data:
                # Add page break for pages after the first
                if page_data['page_number'] > 1:
                    doc.add_page_break()
                
                # Group text blocks into structured content with layout reconstruction
                structured_content = self._group_text_into_paragraphs(page_data['text_blocks'])
                
                # Detect multi-column layout
                columns = self._detect_columns(structured_content)
                
                for content_block in structured_content:
                    text_content = content_block['text']
                    if not text_content.strip():
                        continue
                    
                    content_type = content_block.get('type', 'paragraph')
                    font_size = content_block.get('font_size', 11)
                    font_weight = content_block.get('font_weight', 'normal')
                    
                    # Create appropriate Word element based on content type with better formatting
                    if content_type == 'title':
                        heading = doc.add_heading(text_content, level=0)
                        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # Make title larger and bold
                        for run in heading.runs:
                            run.font.size = Pt(max(16, font_size))
                            run.font.bold = True
                            
                    elif content_type == 'heading':
                        heading = doc.add_heading(text_content, level=1)
                        # Preserve original heading size
                        for run in heading.runs:
                            run.font.size = Pt(max(12, font_size))
                            if font_weight == 'bold':
                                run.font.bold = True
                                
                    else:
                        # Regular paragraph or caption
                        p = doc.add_paragraph()
                        
                        # Add the text with proper formatting
                        run = p.add_run(text_content)
                        run.font.size = Pt(max(8, min(18, font_size)))
                        
                        if font_weight == 'bold':
                            run.font.bold = True
                        
                        # Special formatting for captions
                        if content_type == 'caption':
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run.font.italic = True
                            run.font.size = Pt(max(8, font_size - 1))  # Slightly smaller for captions
                        
                        # Adjust line spacing based on original layout
                        if content_type == 'paragraph':
                            # Set line spacing closer to original
                            p.paragraph_format.line_spacing = 1.15
                            p.paragraph_format.space_after = Pt(6)
            
            doc.save(temp_docx_path)
            print(f"INFO: Created temporary Word document: {temp_docx_path}", file=sys.stderr)
            
            # Convert Word document to PDF using LibreOffice (if available)
            # This creates a text-based PDF that pdf2docx can handle well
            temp_pdf_path = os.path.join(temp_dir, "ocr_temp.pdf")
            
            try:
                import subprocess
                # Try to convert using LibreOffice
                cmd = [
                    'soffice', '--headless', '--convert-to', 'pdf', 
                    '--outdir', temp_dir, temp_docx_path
                ]
                result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
                
                if result.returncode == 0 and os.path.exists(temp_pdf_path):
                    print(f"INFO: Created temporary PDF from OCR: {temp_pdf_path}", file=sys.stderr)
                    return temp_pdf_path
                else:
                    print(f"WARNING: LibreOffice conversion failed: {result.stderr}", file=sys.stderr)
                    # Fall back to returning the Word document path
                    return temp_docx_path
                    
            except Exception as e:
                print(f"WARNING: Could not create PDF from OCR Word document: {e}", file=sys.stderr)
                # Return the Word document path as fallback
                return temp_docx_path
                
        except Exception as e:
            print(f"ERROR: Failed to create temporary document from OCR: {e}", file=sys.stderr)
            return ""

    def _estimate_font_size_accurate(self, width: int, height: int, text_length: int, data: Dict, index: int) -> int:
        """More accurate font size estimation using Tesseract data"""
        if text_length == 0:
            return 11
        
        # Use height as primary indicator (more reliable than width)
        # Most fonts have height roughly equal to point size
        base_size = max(6, min(72, int(height * 0.8)))
        
        # Check if this is part of a larger text block for context
        block_num = data['block_num'][index]
        par_num = data['par_num'][index]
        
        # Find other text in the same paragraph for size consistency
        same_par_heights = []
        for i, (bn, pn, h) in enumerate(zip(data['block_num'], data['par_num'], data['height'])):
            if bn == block_num and pn == par_num and data['text'][i].strip() and int(data['conf'][i]) > 20:
                same_par_heights.append(h)
        
        if same_par_heights:
            # Use median height of paragraph for consistency
            median_height = np.median(same_par_heights)
            base_size = max(6, min(72, int(median_height * 0.8)))
        
        # Adjust based on text characteristics
        text = data['text'][index].strip()
        if text.isupper() and text_length < 20:
            base_size = max(base_size, 14)  # Likely heading
        elif text_length > 100:
            base_size = min(base_size, 12)  # Likely body text
            
        return base_size
    
    def _estimate_font_size(self, width: int, height: int, text_length: int) -> int:
        """Fallback font size estimation for backward compatibility"""
        if text_length == 0:
            return 11
        
        # Approximate font size based on height and character width
        estimated_size = max(8, min(72, int(height * 0.75)))
        
        # Adjust for very short or very long text
        if text_length < 3:
            estimated_size = max(estimated_size, 14)  # Likely a title or heading
        elif text_length > 50:
            estimated_size = min(estimated_size, 12)  # Likely body text
            
        return estimated_size
    
    def _detect_font_weight_improved(self, image: np.ndarray, x: int, y: int, w: int, h: int) -> str:
        """Improved font weight detection using multiple methods"""
        try:
            # Ensure coordinates are within image bounds
            img_h, img_w = image.shape
            x = max(0, min(x, img_w - 1))
            y = max(0, min(y, img_h - 1))
            w = min(w, img_w - x)
            h = min(h, img_h - y)
            
            if w <= 0 or h <= 0:
                return 'normal'
            
            # Extract the text region
            text_region = image[y:y+h, x:x+w]
            
            if text_region.size == 0:
                return 'normal'
            
            # Method 1: Pixel density analysis
            dark_pixels = np.sum(text_region < 128)
            total_pixels = text_region.size
            density = dark_pixels / total_pixels if total_pixels > 0 else 0
            
            # Method 2: Edge density (bold text has thicker strokes)
            edges = cv2.Canny(text_region, 50, 150)
            edge_density = np.sum(edges > 0) / edges.size if edges.size > 0 else 0
            
            # Method 3: Stroke width analysis using morphological operations
            kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (2, 2))
            dilated = cv2.dilate(text_region, kernel, iterations=1)
            stroke_diff = np.sum(dilated != text_region) / text_region.size if text_region.size > 0 else 0
            
            # Combine all methods for better accuracy
            bold_score = (density * 0.5) + (edge_density * 0.3) + (stroke_diff * 0.2)
            
            return 'bold' if bold_score > 0.25 else 'normal'
            
        except Exception:
            pass
        
        return 'normal'
    
    def _detect_font_weight(self, image: np.ndarray, x: int, y: int, w: int, h: int) -> str:
        """Fallback font weight detection for backward compatibility"""
        try:
            # Extract the text region
            text_region = image[y:y+h, x:x+w]
            
            # Calculate pixel density (ratio of dark pixels)
            if text_region.size > 0:
                dark_pixels = np.sum(text_region < 128)
                total_pixels = text_region.size
                density = dark_pixels / total_pixels if total_pixels > 0 else 0
                
                # Higher density usually indicates bold text
                return 'bold' if density > 0.3 else 'normal'
            
        except Exception:
            pass
        
        return 'normal'
    
    def _classify_text_type_improved(self, text: str, font_size: int, y_position: int, data: Dict, index: int) -> str:
        """Improved text type classification using Tesseract structure data"""
        text_lower = text.lower().strip()
        
        # Get structural information from Tesseract
        block_num = data['block_num'][index]
        par_num = data['par_num'][index]
        line_num = data['line_num'][index]
        
        # Check if this is the first line of a paragraph (potential heading)
        is_first_line = True
        for i, (bn, pn, ln) in enumerate(zip(data['block_num'], data['par_num'], data['line_num'])):
            if bn == block_num and pn == par_num and ln < line_num and data['text'][i].strip():
                is_first_line = False
                break
        
        # Check if paragraph has multiple lines (less likely to be heading)
        par_line_count = sum(1 for bn, pn in zip(data['block_num'], data['par_num']) 
                           if bn == block_num and pn == par_num)
        
        # More sophisticated classification
        # Title indicators (usually larger, centered, at top)
        if (font_size > 18 or 
            (y_position < 150 and font_size > 14 and len(text) < 80) or
            (text.isupper() and len(text) < 50 and font_size > 12)):
            return 'title'
        
        # Heading indicators
        if (font_size > 14 or
            (is_first_line and par_line_count == 1 and font_size > 11) or
            any(word in text_lower for word in ['chapter', 'section', 'part', 'introduction', 'conclusion']) or
            (len(text) < 100 and text.endswith(':') and font_size > 10)):
            return 'heading'
        
        # Caption indicators (usually smaller, contains specific keywords)
        if (font_size < 9 or
            any(word in text_lower for word in ['figure', 'table', 'image', 'photo', 'chart', 'graph', 'source']) or
            (len(text) < 30 and any(char.isdigit() for char in text))):
            return 'caption'
        
        return 'paragraph'
    
    def _classify_text_type(self, text: str, font_size: int, y_position: int) -> str:
        """Fallback text type classification for backward compatibility"""
        text_lower = text.lower().strip()
        
        # Check for title indicators
        if (font_size > 16 or 
            (y_position < 100 and len(text) < 50) or
            text.isupper() and len(text) < 30):
            return 'title'
        
        # Check for heading indicators
        if (font_size > 13 or
            any(word in text_lower for word in ['chapter', 'section', 'part']) or
            (len(text) < 80 and text.endswith(':'))):
            return 'heading'
        
        # Check for caption indicators
        if (font_size < 10 or
            any(word in text_lower for word in ['figure', 'table', 'image', 'photo']) or
            len(text) < 20):
            return 'caption'
        
        return 'paragraph'
    
    def _group_text_into_paragraphs(self, text_blocks: List[Dict]) -> List[Dict]:
        """Advanced layout reconstruction using Tesseract structure data"""
        if not text_blocks:
            return []
        
        # Group by Tesseract's structural hierarchy: block -> paragraph -> line -> word
        structured_content = []
        
        # Group blocks by their structural IDs
        blocks_by_structure = {}
        for block in text_blocks:
            block_id = block.get('block_num', 0)
            par_id = block.get('par_num', 0)
            line_id = block.get('line_num', 0)
            
            key = (block_id, par_id, line_id)
            if key not in blocks_by_structure:
                blocks_by_structure[key] = []
            blocks_by_structure[key].append(block)
        
        # Process each line and group into paragraphs
        paragraphs_by_par = {}
        for (block_id, par_id, line_id), line_blocks in blocks_by_structure.items():
            # Sort words in line by horizontal position
            line_blocks.sort(key=lambda x: x['bbox'][0])
            
            # Join words in line with proper spacing
            line_text = []
            last_x_end = 0
            
            for block in line_blocks:
                text = block['text'].strip()
                if not text:
                    continue
                
                bbox = block['bbox']
                x_start = bbox[0]
                
                # Add extra space if there's a significant gap between words
                if last_x_end > 0 and x_start - last_x_end > 20:
                    line_text.append('  ')  # Double space for larger gaps
                elif last_x_end > 0 and x_start - last_x_end > 10:
                    line_text.append(' ')   # Single space for normal gaps
                
                line_text.append(text)
                last_x_end = bbox[2]  # Right edge of current word
            
            if line_text:
                par_key = (block_id, par_id)
                if par_key not in paragraphs_by_par:
                    paragraphs_by_par[par_key] = {
                        'lines': [],
                        'y_positions': [],
                        'font_sizes': [],
                        'font_weights': [],
                        'text_types': []
                    }
                
                paragraphs_by_par[par_key]['lines'].append(''.join(line_text))
                paragraphs_by_par[par_key]['y_positions'].append(line_blocks[0]['bbox'][1])
                paragraphs_by_par[par_key]['font_sizes'].extend([b.get('font_size', 11) for b in line_blocks])
                paragraphs_by_par[par_key]['font_weights'].extend([b.get('font_weight', 'normal') for b in line_blocks])
                paragraphs_by_par[par_key]['text_types'].extend([b.get('text_type', 'paragraph') for b in line_blocks])
        
        # Convert to final structured content
        for (block_id, par_id), par_data in sorted(paragraphs_by_par.items()):
            if not par_data['lines']:
                continue
            
            # Determine paragraph properties
            most_common_font_size = max(set(par_data['font_sizes']), key=par_data['font_sizes'].count) if par_data['font_sizes'] else 11
            most_common_font_weight = max(set(par_data['font_weights']), key=par_data['font_weights'].count) if par_data['font_weights'] else 'normal'
            most_common_text_type = max(set(par_data['text_types']), key=par_data['text_types'].count) if par_data['text_types'] else 'paragraph'
            
            # Join lines with proper line breaks for paragraphs, but not for headings/titles
            if most_common_text_type in ['title', 'heading']:
                paragraph_text = ' '.join(par_data['lines'])
            else:
                # For regular paragraphs, preserve line structure but join sensibly
                paragraph_text = ' '.join(line.strip() for line in par_data['lines'] if line.strip())
            
            if paragraph_text.strip():
                structured_content.append({
                    'text': paragraph_text,
                    'type': most_common_text_type,
                    'font_size': most_common_font_size,
                    'font_weight': most_common_font_weight,
                    'y_position': min(par_data['y_positions']) if par_data['y_positions'] else 0,
                    'block_id': block_id,
                    'par_id': par_id
                })
        
        # Sort by vertical position (top to bottom)
        structured_content.sort(key=lambda x: x['y_position'])
        
        return structured_content
    
    def _detect_columns(self, structured_content: List[Dict]) -> List[List[Dict]]:
        """Detect multi-column layout and group content accordingly"""
        if not structured_content:
            return []
        
        # Analyze horizontal positions to detect columns
        x_positions = []
        for content in structured_content:
            if 'block_id' in content:
                # Use block position data if available
                x_positions.append(content.get('x_position', 0))
        
        if not x_positions:
            return [structured_content]  # Single column
        
        # Use clustering to identify column boundaries
        # Simple approach: if there are distinct horizontal groupings
        x_positions = sorted(set(x_positions))
        
        if len(x_positions) <= 1:
            return [structured_content]  # Single column
        
        # Check for significant gaps that indicate column breaks
        gaps = []
        for i in range(1, len(x_positions)):
            gaps.append(x_positions[i] - x_positions[i-1])
        
        # If there's a large gap, it's likely a column break
        avg_gap = np.mean(gaps) if gaps else 0
        large_gaps = [i for i, gap in enumerate(gaps) if gap > avg_gap * 2]
        
        if large_gaps:
            # Multiple columns detected
            columns = []
            # For now, return single column but this can be enhanced
            # to properly handle multi-column layouts in Word
            return [structured_content]
        
        return [structured_content]  # Single column
    
    def _post_process_ocr_results(self, pages_data: List[Dict]) -> List[Dict]:
        """Post-process OCR results to improve text accuracy"""
        processed_pages = []
        
        for page_data in pages_data:
            processed_blocks = []
            
            for block in page_data['text_blocks']:
                text = block['text']
                
                # Common OCR error corrections
                corrected_text = self._correct_common_ocr_errors(text)
                
                # Update the block with corrected text
                corrected_block = block.copy()
                corrected_block['text'] = corrected_text
                processed_blocks.append(corrected_block)
            
            # Update page data
            processed_page = page_data.copy()
            processed_page['text_blocks'] = processed_blocks
            processed_page['full_text'] = ' '.join([block['text'] for block in processed_blocks if block['text'].strip()])
            processed_pages.append(processed_page)
        
        return processed_pages
    
    def _correct_common_ocr_errors(self, text: str) -> str:
        """Correct common OCR errors to improve text accuracy"""
        if not text:
            return text
        
        # Common character misrecognitions
        corrections = {
            # Number/letter confusions
            'rn': 'm',
            'vv': 'w',
            'VV': 'W',
            
            # Punctuation corrections
            '|': 'l',
            '¡': 'i',
            '§': 's',
            
            # Common word corrections
            'teh': 'the',
            'adn': 'and',
            'hte': 'the',
            'taht': 'that',
            'wihch': 'which',
            'recieve': 'receive',
            'seperate': 'separate',
            'occured': 'occurred',
        }
        
        corrected = text
        
        # Apply basic corrections
        for error, correction in corrections.items():
            corrected = corrected.replace(error, correction)
        
        # Context-sensitive corrections
        # Fix 0/O confusion based on context
        words = corrected.split()
        for i, word in enumerate(words):
            # If word contains digits, 0 is likely correct
            # If word is all letters, 0 should probably be O
            if '0' in word and not any(c.isdigit() for c in word.replace('0', '')):
                words[i] = word.replace('0', 'O')
        
        corrected = ' '.join(words)
        
        # Fix common spacing issues
        corrected = ' '.join(corrected.split())  # Normalize whitespace
        
        return corrected


def main():
    parser = argparse.ArgumentParser(description='Convert PDF to Word using pdf2docx or OCR + pdf2docx')
    parser.add_argument('input_pdf', help='Input PDF file path')
    parser.add_argument('output_file', help='Output Word file path')
    parser.add_argument('--is-scanned', action='store_true', default=False,
                       help='Force OCR mode for scanned PDFs')
    parser.add_argument('--dpi', type=int, default=300, help='DPI for PDF to image conversion (OCR mode)')
    
    args = parser.parse_args()
    
    try:
        # Initialize converter
        converter = PDFToWordConverter()
        
        # Convert based on PDF type
        success = False
        
        if args.is_scanned:
            print("INFO: Using OCR mode for scanned PDF", file=sys.stderr)
            success = converter.convert_scanned_pdf(args.input_pdf, args.output_file, args.dpi)
        else:
            print("INFO: Using pdf2docx for text-based PDF", file=sys.stderr)
            success = converter.convert_text_based_pdf(args.input_pdf, args.output_file)
        
        if not success:
            raise Exception("PDF to Word conversion failed")
        
        # Return success info as JSON
        result = {
            'success': True,
            'output_file': args.output_file,
            'method': 'OCR + pdf2docx' if args.is_scanned else 'pdf2docx',
            'message': 'PDF successfully converted to Word document'
        }
        
        print(f"\nSUCCESS: {json.dumps(result)}")
        
    except Exception as e:
        error_result = {
            'success': False,
            'error': str(e),
            'traceback': traceback.format_exc()
        }
        print(f"\nERROR: {json.dumps(error_result)}")
        sys.exit(1)


if __name__ == '__main__':
    main()
